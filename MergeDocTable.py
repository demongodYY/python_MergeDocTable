 #coding=utf-8
import os
import re
import win32com
from win32com.client import Dispatch, constants
from docx import Document
import sys
reload(sys)
sys.setdefaultencoding('gb2312') #gb2312,gbk 
true = True
false = False

def writeFile(s):
    path = "c:\\test\\result"
    if not os.path.exists(path):
          os.makedirs(path)
    f=file(path + "\\result.csv","a")
    f.writelines(s)
    f.close()

def remove_control_chars(s):
    control_chars = ''.join(map(unichr, range(0,32) + range(127,160)))
    control_char_re = re.compile('[%s]' % re.escape(control_chars))

    return control_char_re.sub('', s)


def parse_doc(f):
  doc = w.Documents.Open( FileName = f )
  count=doc.Paragraphs[1]
  # print (str(count))
  for t in doc.Tables:       
        for row in t.Rows:  
              rowString=""
              for cell in row.Cells:                  
                   rowString = rowString + cell.Range.Text+","
              writeFile(remove_control_chars(str(count)+","+rowString)+"\r\n")
  doc.Close()
def parse_docx(f):
  d = Document(f)
  count=d.paragraphs[1].text
  for t in d.tables:
        for row in t.rows: 
          rowString=""
          for cell in row.cells:              
                rowString = rowString + cell.text+","
          writeFile(remove_control_chars(str(count)+","+rowString)+"\r\n")
                

if __name__ == "__main__":
  w = win32com.client.Dispatch('Word.Application')
  PATH = "c:\\test"
  doc_files = os.listdir(PATH)
  for doc in doc_files:
    if os.path.splitext(doc)[1] == '.docx' and doc[0]!='~':
      try:
        parse_docx(PATH+'\\'+doc)
      except Exception as e:
         print (e)
    elif os.path.splitext(doc)[1] == '.doc' and doc[0]!='~':
      try:
        parse_doc(PATH+'\\'+doc)
      except Exception as e:
         print (e)
